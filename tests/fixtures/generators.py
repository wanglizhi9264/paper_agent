"""Generate small, deterministic PDF and DOCX fixtures for golden tests.

PDF generation runs in a **subprocess** because PyMuPDF's native code segfaults
inside the pytest process on some macOS/arm64 environments. The generated
files are tiny, clearly licensed (MIT, no third-party content), and live in
the session-scoped tmp path.
"""

from __future__ import annotations

import subprocess
import sys
from pathlib import Path

_GEN_SCRIPT = """
import sys
from pathlib import Path

mode = sys.argv[1]
out = Path(sys.argv[2])

if mode == "pdf":
    import pymupdf
    doc = pymupdf.open()
    p1 = doc.new_page()
    p1.insert_text((72, 72), "Sample Paper Title", fontsize=16)
    p1.insert_text((72, 110), "Introduction", fontsize=13)
    p1.insert_text((72, 140), "This is a sample paragraph for testing the PDF loader.", fontsize=11)
    p2 = doc.new_page()
    p2.insert_text((72, 72), "Method", fontsize=13)
    p2.insert_text((72, 100), "The method section describes the approach used in the paper.", fontsize=11)
    doc.save(str(out))
    doc.close()

elif mode == "ocr_pdf":
    import pymupdf
    doc = pymupdf.open()
    for _ in range(5):
        doc.new_page()  # completely empty pages
    doc.save(str(out))
    doc.close()

elif mode == "docx":
    import docx
    document = docx.Document()
    document.add_heading("Sample DOCX Title", level=0)
    document.add_heading("Introduction", level=1)
    document.add_paragraph("This is a sample paragraph in the DOCX fixture.")
    document.add_heading("Data", level=2)
    table = document.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    table.cell(0, 0).text = "Metric"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Accuracy"
    table.cell(1, 1).text = "0.95"
    document.save(str(out))
"""


def _run(mode: str, path: Path) -> Path:
    """Generate a fixture file by running the generator script in a subprocess."""
    path.parent.mkdir(parents=True, exist_ok=True)
    result = subprocess.run(
        [sys.executable, "-c", _GEN_SCRIPT, mode, str(path)],
        capture_output=True,
        text=True,
        check=False,
    )
    if result.returncode != 0:
        raise RuntimeError(f"fixture generation '{mode}' failed: {result.stderr or result.stdout}")
    return path


def make_pdf(path: Path) -> Path:
    return _run("pdf", path)


def make_docx(path: Path) -> Path:
    return _run("docx", path)


def make_ocr_pdf(path: Path) -> Path:
    return _run("ocr_pdf", path)
